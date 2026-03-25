#!/usr/bin/env python3
"""Generate Mela Co-Founder Growth Playbook as a professional .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── colour palette ────────────────────────────────────────────────────────────
ORANGE   = RGBColor(0xE8, 0x76, 0x0A)   # #E8760A
DARK     = RGBColor(0x1A, 0x1A, 0x1A)   # #1A1A1A
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
GREY_BG  = RGBColor(0xF2, 0xF2, 0xF2)   # light grey for code blocks
CREAM    = RGBColor(0xFA, 0xFA, 0xF7)

BODY_FONT  = "Calibri"
CODE_FONT  = "Courier New"


# ── helper: set cell background colour ───────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    """Add thin borders to every cell in the table."""
    tbl  = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tblBorders.append(el)
    tblPr.append(tblBorders)


# ── helper: add a styled paragraph ───────────────────────────────────────────
def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(18)
    run.font.bold  = True
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # orange bottom border via shading on paragraph — use a rule line instead
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E8760A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name   = BODY_FONT
    run.font.size   = Pt(11)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_numbered(doc, text, number):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_code_block(doc, text):
    """Grey-background monospace block."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    # grey shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent   = Inches(0.2)
    p.paragraph_format.right_indent  = Inches(0.2)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Create a bordered table with orange header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_borders(table)

    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "E8760A")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name  = BODY_FONT
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = WHITE

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "FAFAF7"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name  = BODY_FONT
            run.font.size  = Pt(10)
            run.font.color.rgb = DARK

    # column widths
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # spacer
    return table


def add_page_break(doc):
    doc.add_page_break()


def add_spacer(doc, size_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size_pt)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT BUILD
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# ── footer with page numbers ──────────────────────────────────────────────────
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run("Mela — Co-Founder Growth & Brand Playbook  |  melaa.ca  |  Page ").font.size = Pt(9)
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
run_pg = fp.add_run()
run_pg.font.size = Pt(9)
run_pg._r.append(fldChar1)
run_pg._r.append(instrText)
run_pg._r.append(fldChar2)


# ═══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
add_spacer(doc, 60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mela")
run.font.name  = BODY_FONT
run.font.size  = Pt(42)
run.font.bold  = True
run.font.color.rgb = ORANGE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Co-Founder Growth & Brand Playbook")
run2.font.name  = BODY_FONT
run2.font.size  = Pt(26)
run2.font.bold  = True
run2.font.color.rgb = DARK

add_spacer(doc, 10)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Your complete step-by-step guide to launching the brand,\nbuilding the community, and getting the first 50 vendors")
run3.font.name   = BODY_FONT
run3.font.size   = Pt(13)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_spacer(doc, 40)

for line, sz, bold in [
    ("March 2026", 12, False),
    ("Prepared by: Ishaan Saini, Founder", 12, True),
    ("melaa.ca", 12, False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(sz)
    run.font.bold  = bold
    run.font.color.rgb = DARK

add_spacer(doc, 50)

# note box
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_note._p.get_or_add_pPr()
shd = OxmlElement("w:shd")
shd.set(qn("w:val"),   "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"),  "FFF3E0")
pPr.append(shd)
run_note = p_note.add_run(
    '"The AI agents are running 24/7 in the background handling emails, outreach, '
    'and revenue monitoring. Your job is everything human — brand, community, relationships."'
)
run_note.font.name   = BODY_FONT
run_note.font.size   = Pt(11)
run_note.font.italic = True
run_note.font.color.rgb = ORANGE
p_note.paragraph_format.left_indent  = Inches(0.5)
p_note.paragraph_format.right_indent = Inches(0.5)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 1: THE MISSION
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 1: The Mission")
add_heading2(doc, "What We're Building & Why It Matters")

add_body(doc, "Mela (melaa.ca) is the GTA's only marketplace built exclusively for South Asian weddings.")
add_body(doc, "The problem we solve: South Asian couples in Toronto, Brampton, and Mississauga spend hours searching through generic directories like WeddingWire and Google — and most results don't understand the culture, the traditions, or what makes a South Asian wedding what it is.")
add_body(doc, "The market: South Asian weddings in Canada are a $2B+ industry. The GTA has the highest concentration of South Asian families in North America. No one has built a dedicated marketplace for this community. We are.")

add_heading3(doc, "Current State — March 2026")
for item in [
    "Platform: Live at melaa.ca",
    "Active Vendors: 5 verified",
    "MRR: $795/month",
    "Valuation: $76,320",
    "AI Agents: 9 departments running 24/7, automatically sending emails, generating content, and monitoring revenue",
]:
    add_bullet(doc, item)

add_heading3(doc, "Your Role: Growth & Community Co-Founder")
add_body(doc, "You own: Instagram, TikTok, Facebook, community building, vendor relationships, brand voice, content creation, and all human-touch outreach.")
add_body(doc, "The agents handle: Cold email sequences, lead notifications, SEO blog drafts, revenue monitoring, and financial reporting — automatically, around the clock.")

add_heading3(doc, "90-Day Goals")
for item in [
    "50+ active vendors listed",
    "$2,000 MRR",
    "1,000 Instagram followers",
    "500 email subscribers",
]:
    add_bullet(doc, item)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 2: WEEK 1 PRIORITY TIMELINE
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 2: Week 1 Priority Timeline")
add_heading2(doc, "Week 1: Foundation (Days 1–7)")
add_body(doc, "This is your most important week. Everything you do in Week 1 sets the foundation for the next 90 days. Follow this day by day.")

day_data = [
    ("Day 1 — Brand Setup (4–5 hours)", [
        "Create Instagram account (@melaa.ca)",
        "Create TikTok account (@melaa.ca)",
        "Create Facebook Business Page",
        "Create Pinterest account",
        "Create LinkedIn Company Page",
        "Set up Canva account with brand colours",
        "Post Day 1 Instagram content (copy-paste in Section 5)",
        "Email 15 people from your personal network (copy-paste in Section 10)",
        "Post in 3 Facebook groups (copy-paste in Section 7)",
    ]),
    ("Day 2 — Vendor Outreach (3 hours)", [
        "Go to melaa.ca/admin/outreach — review all 20 outreach targets",
        "WhatsApp 10 of those vendors (script in Section 8)",
        "DM 10 vendors on Instagram using hashtag search (script in Section 9)",
        "Set up Google Business Profile (step-by-step in Section 11)",
    ]),
    ("Day 3 — Content & Community (2 hours)", [
        "Post 3 Reddit threads (copy-paste in Section 12)",
        "Film and post first TikTok video (script in Section 6)",
        "List Mela on WeddingWire Canada and The Knot (steps in Section 13)",
        "Follow 50 South Asian wedding vendors on Instagram",
    ]),
    ("Day 4 — Community Deep Dive (2 hours)", [
        "Join all 10 Facebook groups listed in Section 7",
        "Post in 5 groups using the copy-paste scripts",
        "Search and follow 50 more South Asian vendors on Instagram",
        "Leave genuine comments on 20 vendor posts",
    ]),
    ("Day 5 — Partnership Outreach (2 hours)", [
        "DM 5 South Asian wedding influencers (script in Section 15)",
        "Email 3 South Asian wedding blogs (script in Section 15)",
        "Continue vendor DM outreach — aim for 10 more contacts",
    ]),
    ("Day 6 — Content Production (3 hours)", [
        "Film 2–3 TikTok videos (scripts in Section 6)",
        "Create 5 Instagram posts in Canva using scripts in Section 5",
        "Post LinkedIn announcement (copy-paste in Section 14)",
    ]),
    ("Day 7 — Review & Plan (1 hour)", [
        "Check melaa.ca/admin for new vendor signups",
        "Review which outreach got responses",
        "Follow up all unanswered DMs from Week 1",
        "Plan Week 2 content",
    ]),
]

for day_title, steps in day_data:
    add_heading3(doc, day_title)
    for i, step in enumerate(steps, 1):
        add_numbered(doc, step, i)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 3: 90-DAY GROWTH ROADMAP
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 3: 90-Day Growth Roadmap")
add_heading2(doc, "The 90-Day Brand & Growth Strategy")

add_table(
    doc,
    headers=["Month", "Goals", "Focus Areas", "Key Actions"],
    rows=[
        [
            "Month 1\n(Days 1–30)",
            "20+ vendors, $1,200 MRR, 300 IG followers",
            "Instagram daily, personal network, Facebook groups, WhatsApp outreach",
            "Post daily, WhatsApp 10 vendors/day, post in 2 groups/day, personal emails",
        ],
        [
            "Month 2\n(Days 31–60)",
            "35 vendors, $1,600 MRR, 700 IG followers",
            "TikTok growth, influencer partnerships, vendor success stories",
            "3 TikToks/week, 2 influencer collabs, collect vendor testimonials",
        ],
        [
            "Month 3\n(Days 61–90)",
            "50 vendors, $2,000 MRR, 1,000 IG followers",
            "Press coverage, community events, partnership deals",
            "PR pitch to 5 outlets, host one community event, YouTube launch",
        ],
    ],
    col_widths=[1.2, 1.6, 1.8, 2.0],
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 4: INSTAGRAM SETUP
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 4: Instagram Setup — Complete Guide")
add_heading2(doc, "Setting Up @melaa.ca on Instagram")

add_heading3(doc, "Step 1: Create the Account")
for i, step in enumerate([
    "Go to instagram.com on your phone",
    "Tap \"Sign Up\"",
    "Use email: social@melaa.ca (ask Ishaan to create this)",
    "Username: try melaa.ca first, then mela.weddings, then melaa_weddings",
    "Switch to Professional Account → Business → Wedding Services",
], 1):
    add_numbered(doc, step, i)

add_heading3(doc, "Step 2: Profile — Copy and Paste Exactly")

add_body(doc, "Profile name — copy this exactly:")
add_code_block(doc, "Mela | South Asian Weddings GTA")

add_body(doc, "Bio — copy this exactly:")
add_code_block(doc,
    "🌺 GTA's South Asian wedding vendor marketplace\n"
    "📍 Toronto · Brampton · Mississauga · Markham\n"
    "✨ Find verified photographers, caterers & more\n"
    "🔗 melaa.ca"
)

add_heading3(doc, "Step 3: Canva Brand Kit")
add_body(doc, "Create a free Canva account at canva.com and set up these brand colours:")

add_table(
    doc,
    headers=["Colour", "Hex Code"],
    rows=[
        ["Primary Orange", "#E8760A"],
        ["Dark / Almost Black", "#1A1A1A"],
        ["Cream Background", "#FAFAF7"],
        ["Light Grey Border", "#E5E5E0"],
    ],
    col_widths=[3.0, 3.0],
)

add_body(doc, "Fonts: Playfair Display for headings, Inter or DM Sans for body text.")

add_heading3(doc, "Step 4: Create Story Highlights")
add_body(doc, "Create these 5 Highlights immediately (use orange covers in Canva):")
for item in ["Our Vendors", "Real Weddings", "How It Works", "List Free", "Testimonials"]:
    add_bullet(doc, item)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 5: INSTAGRAM CONTENT
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 5: Instagram Content — First 30 Posts")
add_heading2(doc, "Instagram Content: First 30 Posts (Copy-Paste Ready)")
add_body(doc, "Rotate between 3 content types: (1) Vendor Spotlights, (2) Wedding Tips for couples, (3) Community & Culture posts.")

posts = [
    (
        "Post 1 — Launch Day (Day 1)",
        "Visual: Orange background, white text \"Mela — South Asian Weddings GTA\". Create in Canva.",
        "Caption — copy and paste:",
        "We built this for you 🧡\n\n"
        "Finding the right vendors for your South Asian wedding in the GTA shouldn't mean spending hours on Google getting results that don't understand your culture.\n\n"
        "Mela is the GTA's first marketplace built exclusively for South Asian weddings. Every vendor understands your traditions, your language, and what makes your celebration uniquely yours.\n\n"
        "Photographers 📸 | Decorators 🌸 | Caterers 🍽️ | Mehndi 🤍 | DJs 🎵\n\n"
        "Free to browse. Free for vendors to list.\n\n"
        "Link in bio → melaa.ca\n\n"
        "#SouthAsianWedding #GTAWedding #SouthAsianBride #DesiWedding #BramptonWeddings #MississaugaWeddings #TorontoWeddings #WeddingVendors #IndianWedding #PunjabiWedding"
    ),
    (
        "Post 2 — Problem/Solution (Day 2)",
        "Visual: Two-panel image — left shows generic Google results, right shows Mela. Create in Canva with orange and dark theme.",
        "Caption — copy and paste:",
        "The struggle was real 😅\n\n"
        "Searching for a South Asian wedding photographer in Brampton and getting results for photographers who've never shot a baraat in their life.\n\n"
        "That's why we built Melaa.ca 🧡\n\n"
        "Every single vendor on Mela:\n"
        "✅ Is local to the GTA\n"
        "✅ Understands South Asian culture and traditions\n"
        "✅ Has been reviewed before listing\n\n"
        "Browse free at melaa.ca (link in bio)\n\n"
        "#SouthAsianWedding #DesiWedding #GTABride #BramptonWedding #WeddingPlanning #WeddingVendors #SouthAsianBride"
    ),
    (
        "Post 3 — Vendor Spotlight (use for every new vendor)",
        "Visual: Photo of the vendor's work (ask them for a photo), add Mela logo watermark in corner.",
        "Caption template — fill in the blanks:",
        "Vendor Spotlight 🌟\n\n"
        "Meet [Vendor Name] — [Category] based in [City], GTA\n\n"
        "[One sentence about what makes them special — look at their Instagram]\n\n"
        "[Vendor Name] is now listed on Melaa.ca and taking bookings for 2026–2027 weddings.\n\n"
        "Find them and other verified South Asian vendors at melaa.ca (link in bio) 🧡\n\n"
        "#SouthAsianWedding #[City]Weddings #WeddingVendor #GTAWeddings #DesiWedding"
    ),
    (
        "Post 4 — Wedding Tip (Day 4)",
        None,
        "Caption — copy and paste:",
        "3 questions you MUST ask your caterer before booking 📝\n\n"
        "1️⃣ Have you catered South Asian weddings before? How many?\n\n"
        "2️⃣ Can you accommodate both vegetarian and non-veg guests?\n\n"
        "3️⃣ Do you handle setup and breakdown, or is that extra?\n\n"
        "Save this for when you're vendor hunting 🔖\n\n"
        "Find verified South Asian caterers in the GTA at melaa.ca (link in bio)\n\n"
        "#SouthAsianWedding #WeddingTips #DesiWedding #WeddingCaterer #IndianWedding #WeddingPlanning #GTAWeddings"
    ),
    (
        "Post 5 — Vendor CTA (Day 5)",
        None,
        "Caption — copy and paste:",
        "Attention GTA wedding vendors 📣\n\n"
        "If you offer services for South Asian weddings and you're not on Melaa.ca yet, you're missing leads.\n\n"
        "Here's the deal:\n"
        "🆓 Free to list\n"
        "⚡ Profile live same day\n"
        "📩 Leads sent directly to you\n"
        "🔒 Founding rate: $49/mo locked forever (normally $197/mo)\n\n"
        "First 50 vendors get the founding rate.\n\n"
        "Link in bio → melaa.ca/list-your-business\n\n"
        "#WeddingVendor #SouthAsianWedding #GTABusiness #WeddingBusiness #BramptonBusiness #MississaugaBusiness"
    ),
    (
        "Post 6 — Community Story (Day 6)",
        None,
        "Caption — copy and paste:",
        "A wedding vendor marketplace made by the community, for the community 🧡\n\n"
        "We're not a corporation. We're not a big tech company. We're South Asian, we're from the GTA, and we built this because we knew this community deserved something better.\n\n"
        "If you're planning your wedding or you work in the industry, come be part of what we're building.\n\n"
        "melaa.ca — link in bio 🌺\n\n"
        "#SouthAsianWedding #GTACommunity #DesiWedding #MadeForUs #SouthAsianCommunity"
    ),
    (
        "Post 7 — Mehndi Tips (Day 7)",
        None,
        "Caption — copy and paste:",
        "How to choose your mehndi artist 🤍\n\n"
        "Things most brides don't think about until it's too late:\n\n"
        "✦ Ask to see their bridal portfolio specifically — not just regular henna\n\n"
        "✦ Book a trial session 1–2 months before the wedding\n\n"
        "✦ Confirm how many artists they bring for large bridal parties\n\n"
        "✦ Ask about paste quality — natural henna vs chemical mix matters for stain colour\n\n"
        "✦ Confirm: do they come to your venue or do you go to them?\n\n"
        "Find verified mehndi artists in the GTA at melaa.ca 🧡\n\n"
        "#MehndiArtist #BridalMehndi #SouthAsianWedding #DesiWedding #GTAMehndi #WeddingMehndi #IndianBride"
    ),
    (
        "Post 8 — FAQ (use whenever engagement drops)",
        None,
        "Caption — copy and paste:",
        '"Is Mela free?"\n\n'
        "Yes. 100% free to browse and find vendors.\n\n"
        "Free for vendors to create a profile too.\n\n"
        "No hidden fees. No booking commissions.\n\n"
        "The vendor and couple connect directly — we just make the introduction.\n\n"
        "That's it. That's the whole model.\n\n"
        "Check it out → melaa.ca (link in bio) 🧡\n\n"
        "#SouthAsianWedding #WeddingPlanning #DesiWedding #GTAWeddings #WeddingTips"
    ),
    (
        "Post 9 — Photography Tips",
        None,
        "Caption — copy and paste:",
        "What to look for in a South Asian wedding photographer 📸\n\n"
        "✦ Have they shot a baraat before? Can they handle fast-moving ceremony moments?\n\n"
        "✦ Do they have a second shooter for large ceremonies?\n\n"
        "✦ How do they handle low-light mandap photography?\n\n"
        "✦ What's their turnaround time for edited photos?\n\n"
        "✦ Do they offer both photos and video, or just one?\n\n"
        "Find South Asian wedding photographers in your GTA city at melaa.ca 🧡\n\n"
        "#SouthAsianWeddingPhotographer #WeddingPhotography #GTAPhotographer #DesiWedding #BridalPhotography"
    ),
    (
        "Post 10 — Growth/Stats",
        None,
        "Caption — copy and paste:",
        "Something is building 👀\n\n"
        "Couples across the GTA are searching for South Asian wedding vendors every single day.\n\n"
        "Photographers in Brampton 📸\n"
        "Decorators in Mississauga 🌸\n"
        "Caterers in Markham 🍽️\n"
        "Mehndi artists in Toronto 🤍\n\n"
        "They're searching. Are you listed?\n\n"
        "Free to list your business → melaa.ca/list-your-business (link in bio)\n\n"
        "#SouthAsianWedding #WeddingBusiness #GrowYourBusiness #DesiWedding #WeddingVendor"
    ),
]

for title, visual, caption_label, caption in posts:
    add_heading3(doc, title)
    if visual:
        add_body(doc, visual, italic=True)
    add_body(doc, caption_label)
    add_code_block(doc, caption)
    add_spacer(doc, 4)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 6: TIKTOK STRATEGY
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 6: TikTok Strategy")
add_heading2(doc, "TikTok Setup + First 10 Video Scripts")

add_heading3(doc, "Account Setup")
for i, step in enumerate([
    "Download TikTok on your phone",
    "Create account with email: social@melaa.ca",
    "Username: @melaa.ca",
    "Bio — copy and paste exactly:",
], 1):
    add_numbered(doc, step, i)

add_code_block(doc,
    "GTA's South Asian wedding vendor marketplace 🌺\n"
    "Find vetted photographers, decorators & more\n"
    "🔗 melaa.ca"
)

tiktok_videos = [
    (
        "Video 1 — Launch Story (film yourself talking, casual)",
        "Script — copy and paste:",
        '[Hook — say this in the first 2 seconds]\n'
        '"So I built something for the South Asian community in the GTA."\n\n'
        '[Middle]\n'
        '"It\'s called Mela. Basically a marketplace specifically for South Asian wedding vendors in the GTA.\n\n'
        'The reason I built it? Every time someone I know plans a South Asian wedding, they spend weeks Googling photographers and decorators who have never shot a South Asian wedding in their life.\n\n'
        'Mela fixes that. Every single vendor on the platform is local to the GTA and knows South Asian culture."\n\n'
        '[CTA]\n'
        '"It\'s free to browse, free for vendors to list. Check it out at melaa.ca — link in bio."'
    ),
    (
        "Video 2 — POV / Text-on-screen style (use trending audio)",
        "Script — copy and paste:",
        '[Slide 1 - text on screen]\n'
        'POV: You\'re planning your South Asian wedding and searching Google for a photographer in Brampton...\n\n'
        '[Slide 2]\n'
        'Result 1: Swedish photographer 3 hours away 🙃\n\n'
        '[Slide 3]\n'
        'Result 2: "We love all cultures!" (never shot a baraat)\n\n'
        '[Slide 4]\n'
        'Result 3: $15,000 minimum. Based in Vancouver.\n\n'
        '[Slide 5]\n'
        'This is why I built melaa.ca 🧡\n\n'
        '[Slide 6 — show screen recording of website]\n'
        'GTA only. South Asian vendors only. Free to browse.\n'
        'Link in bio.'
    ),
    (
        "Video 3 — Vendor Pitch (talking head)",
        "Script — copy and paste:",
        '[On camera, direct]\n'
        '"If you\'re a wedding vendor in the GTA who works with South Asian weddings, here\'s something you should know.\n\n'
        'Mela is a brand new directory specifically for South Asian wedding vendors. We\'re in growth mode right now, which means the first vendors to list are going to get the most visibility.\n\n'
        'It\'s free to list, your profile goes live same day, and couples are already browsing.\n\n'
        'Go to melaa.ca/list-your-business — link in bio."'
    ),
    (
        "Video 4 — Wedding tip (value content)",
        "Script — copy and paste:",
        '[On camera or text slides]\n'
        '"3 things no one tells you about hiring a mehndi artist for your South Asian wedding:"\n\n'
        '"One — book at least 3 months in advance. Good mehndi artists in the GTA book out fast for wedding season."\n\n'
        '"Two — always ask to see their bridal portfolio specifically, not just their regular henna work."\n\n'
        '"Three — confirm if they come to your venue or if you go to them. It makes a huge difference on wedding day."\n\n'
        '"Find verified mehndi artists in the GTA at melaa.ca — link in bio."'
    ),
]

for title, label, script in tiktok_videos:
    add_heading3(doc, title)
    add_body(doc, label)
    add_code_block(doc, script)
    add_spacer(doc, 4)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 7: FACEBOOK GROUP STRATEGY
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 7: Facebook Group Strategy")
add_heading2(doc, "Facebook Groups — Join + Copy-Paste Posts")

add_heading3(doc, "Groups to Join (search each on Facebook)")
fb_groups = [
    "South Asian Brides Canada",
    "Desi Weddings GTA",
    "Brampton South Asian Community",
    "Mississauga Desi Community",
    "South Asian Wedding Planning Canada",
    "Toronto Desi Weddings",
    "Punjabi Weddings GTA",
    "Pakistani Weddings Canada",
    "Bengali Weddings Toronto",
    "Tamil Weddings GTA",
]
for i, g in enumerate(fb_groups, 1):
    add_numbered(doc, g, i)

fb_posts = [
    (
        "Post 1 — Couple Perspective (most natural)",
        "Copy and paste exactly:",
        "Hey everyone! My friend is planning her South Asian wedding and told me about this site called Melaa.ca that has GTA vendors specifically for South Asian weddings — photographers, decorators, caterers, mehndi artists etc all in one place.\n\n"
        "Has anyone used it or know any vendors listed there? She wants to know if the vendors are good before she reaches out 🙏"
    ),
    (
        "Post 2 — Direct Value for Vendor Groups",
        "Copy and paste exactly:",
        "Heads up for any GTA wedding vendors in this group 👋\n\n"
        "There's a new directory called Melaa.ca specifically for South Asian wedding vendors in the GTA. I just listed my business and already got an inquiry.\n\n"
        "It's free to list and they have a Founding Vendor rate of $49/mo if you want priority placement (normally $197/mo).\n\n"
        "Just sharing in case anyone else finds it useful!\n"
        "melaa.ca/list-your-business"
    ),
    (
        "Post 3 — Planning Help Angle",
        "Copy and paste exactly:",
        "Anyone planning a South Asian wedding in the GTA in 2026 or 2027?\n\n"
        "Drop your city below and I'll try to share some vendor recommendations! I've been using Melaa.ca lately which has a good collection of GTA vendors specifically for South Asian weddings 🧡"
    ),
]

for title, label, content in fb_posts:
    add_heading3(doc, title)
    add_body(doc, label)
    add_code_block(doc, content)
    add_spacer(doc, 4)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 8: WHATSAPP VENDOR OUTREACH
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 8: WhatsApp Vendor Outreach")
add_heading2(doc, "WhatsApp Outreach — Step-by-Step + Scripts")

add_heading3(doc, "How to Find the 20 Outreach Targets")
for i, step in enumerate([
    "Go to melaa.ca/admin/outreach (get login from Ishaan)",
    "You'll see 20 vendor targets already identified with business names, categories, and cities",
    "Google each business name to find their phone number or WhatsApp",
    "Alternatively, find them on Instagram — most vendors have WhatsApp in bio",
], 1):
    add_numbered(doc, step, i)

add_heading3(doc, "Opening Message — copy and paste, fill in the blanks")
add_code_block(doc,
    "Hey [Name] 👋 I'm [Your Name], I work with Melaa.ca — a South Asian wedding vendor directory for the GTA.\n\n"
    "We're getting couples searching for [Photography/Decor/Catering — their category] in [their city] every day and I'd love to feature [Business Name] for free.\n\n"
    "Takes literally 5 minutes to set up your profile. Would this be useful for you?"
)

add_heading3(doc, "If they say YES")
add_code_block(doc,
    "Amazing! Here's the link: melaa.ca/list-your-business\n\n"
    "Just fill out the form and your profile goes live same day. You'll get a notification every time a couple contacts you directly through the platform.\n\n"
    "Let me know if you have any questions — happy to help!"
)

add_heading3(doc, "If they ask about cost")
add_code_block(doc,
    "It's completely free to list! There's also a Founding Vendor upgrade for $49/mo (locked forever, normally $197/mo) that gives you priority placement and a verified badge — but the free listing is fully functional and gets you leads too."
)

add_heading3(doc, "Follow-up if no reply after 3 days")
add_code_block(doc,
    "Hey [Name], just following up in case this got buried!\n\n"
    "Quick question — do you currently get leads from online sources or is it mostly word of mouth? Asking because I want to make sure Mela is actually useful for vendors like you before I reach out further."
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 9: INSTAGRAM DM SCRIPTS
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 9: Instagram DM Scripts")
add_heading2(doc, "Instagram DM Outreach — Vendor Acquisition")

add_heading3(doc, "Hashtags to Search for Vendors")
for tag in [
    "#bramptonphotographer",
    "#mississaugaweddingphotographer",
    "#gtaweddingphotographer",
    "#southasianweddingphotographer",
    "#bramptonweddingdecor",
    "#gtamehndi",
    "#southasianweddingcatering",
    "#mississaugaweddingdj",
    "#gtasouthasianwedding",
    "#torontosouthasianwedding",
    "#bramptonmakeup",
    "#gtaweddingmakeup",
]:
    add_bullet(doc, tag)

add_heading3(doc, "DM Script — Initial Message")
add_code_block(doc,
    "Hey [Name]! Love your work 🙌\n\n"
    "I'm [Your Name] from Mela — we're building a wedding vendor marketplace specifically for South Asian weddings in the GTA (melaa.ca).\n\n"
    "We're in early days and handpicking the best vendors to feature. Would you be open to a free listing? Couples in [their city] are already searching on the platform."
)

add_heading3(doc, "DM Follow-up (Day 3 if no reply)")
add_code_block(doc,
    "Hey [Name], following up on my message! We just had a couple search for [their category] in [their city] today — wanted to make sure you're listed before they find someone else.\n\n"
    "Free to list: melaa.ca/list-your-business 🧡"
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 10: PERSONAL NETWORK EMAIL
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 10: Personal Network Email")
add_heading2(doc, "Personal Network Email — Copy and Paste")

add_body(doc, "Send this to 15–20 people you know personally who are South Asian and aged 22–40. This has the highest conversion rate of any outreach.")

add_body(doc, "Subject line — copy exactly:")
add_code_block(doc, "quick favour — something I built")

add_body(doc, "Email body — copy and paste:")
add_code_block(doc,
    "Hey [Name],\n\n"
    "Built something and wanted to share it with people I actually know before I go full launch mode.\n\n"
    "It's called Mela — melaa.ca — a wedding vendor marketplace built specifically for South Asian weddings in the GTA. Think photographers, caterers, decorators, mehndi artists — all in one place, all local, all familiar with South Asian culture.\n\n"
    "Two things that would mean a lot:\n\n"
    "1. If you know anyone planning a South Asian wedding, send them the link: melaa.ca\n\n"
    "2. If you know any wedding vendors (photographers, decorators, caterers etc), send them this: melaa.ca/list-your-business — it's completely free to list and they'll get leads from couples already searching.\n\n"
    "That's it. No ask beyond that.\n\n"
    "Thanks for any support 🙏\n\n"
    "— [Your Name]"
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 11: GOOGLE BUSINESS PROFILE
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 11: Google Business Profile")
add_heading2(doc, "Google Business Profile — Step-by-Step Setup")

steps_gbp = [
    ("Go to business.google.com on your computer", None),
    ("Click \"Manage now\" or \"Add your business\"", None),
    ("Enter the business name — copy exactly:", "Mela — South Asian Wedding Vendors GTA"),
    ("Category — type and select:", "Wedding service"),
    ("Website — copy exactly:", "https://melaa.ca"),
    ("Service area — add all of these cities:",
     "Toronto, ON\nBrampton, ON\nMississauga, ON\nMarkham, ON\nVaughan, ON\nScarborough, ON\nRichmond Hill, ON\nOakville, ON\nAjax, ON"),
    ("Business description — copy and paste exactly:",
     "Mela is the GTA's only marketplace built exclusively for South Asian weddings. Find verified photographers, decorators, caterers, mehndi artists, DJs, and makeup artists — all serving the South Asian community across the Greater Toronto Area. Free to browse. Free for vendors to list. Couples connect directly with vendors — no commissions, no middlemen."),
    ("Choose phone verification (fastest — you get a code by text)", None),
    ("Once verified, upload photos: screenshots of melaa.ca and the Mela logo", None),
]

for i, (step, code) in enumerate(steps_gbp, 1):
    add_numbered(doc, step, i)
    if code:
        add_code_block(doc, code)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 12: REDDIT POSTS
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 12: Reddit Posts")
add_heading2(doc, "Reddit Posts — Copy and Paste")

reddit_posts = [
    (
        "r/SouthAsianWeddings",
        "Title — copy exactly:",
        "Built a free South Asian wedding vendor directory for GTA — would love honest feedback",
        "Body — copy and paste:",
        "Hey everyone! I built melaa.ca — a directory specifically for South Asian wedding vendors in the Greater Toronto Area.\n\n"
        "The problem I kept hearing: South Asian couples spend hours Googling photographers, decorators, caterers individually and half the results aren't South Asian or don't understand the culture.\n\n"
        "Mela fixes that. Every vendor is local to GTA and understands South Asian weddings. Free to browse, free for vendors to list.\n\n"
        "Would love brutal feedback — what's missing, what's wrong, what would make it actually useful for you?"
    ),
    (
        "r/brampton",
        "Title — copy exactly:",
        "Made a free website to find South Asian wedding vendors in Brampton/GTA — thoughts?",
        "Body — copy and paste:",
        "Hey Brampton! I built melaa.ca — basically a Google specifically for South Asian wedding vendors. Photographers, caterers, mehndi artists, decorators, DJs — all local to GTA, all familiar with South Asian weddings.\n\n"
        "It's completely free to browse and free for vendors to list their business.\n\n"
        "Would love to know if this is something the community would actually use. What am I missing?"
    ),
    (
        "r/toronto",
        "Title — copy exactly:",
        "Free directory for South Asian wedding vendors in GTA (feedback welcome)",
        "Body — copy and paste:",
        "Built melaa.ca — a vendor marketplace specifically for South Asian weddings in the GTA. Tried to solve the problem of spending hours hunting through generic directories that have nothing South Asian specific.\n\n"
        "Completely free. Would love feedback from anyone who's planned or is planning a wedding."
    ),
]

for subreddit, title_label, title_text, body_label, body_text in reddit_posts:
    add_heading3(doc, subreddit)
    add_body(doc, title_label)
    add_code_block(doc, title_text)
    add_body(doc, body_label)
    add_code_block(doc, body_text)
    add_spacer(doc, 6)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 13: WEDDINGWIRE + THE KNOT
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 13: WeddingWire + The Knot")
add_heading2(doc, "List Mela on Wedding Directories")

add_heading3(doc, "WeddingWire Canada")
ww_steps = [
    ("Go to weddingwire.ca", None),
    ("Click \"List your business\"", None),
    ("Business Name — copy exactly:", "Mela — South Asian Wedding Vendors GTA"),
    ("Category: Wedding Planning", None),
    ("Website: melaa.ca", None),
    ("Description — copy and paste:",
     "Mela is the GTA's dedicated marketplace for South Asian wedding vendors. Browse verified photographers, decorators, caterers, mehndi artists, DJs and more — all serving the South Asian community across Toronto, Brampton, Mississauga, Markham, and Vaughan. Free to browse, free to list."),
]
for i, (step, code) in enumerate(ww_steps, 1):
    add_numbered(doc, step, i)
    if code:
        add_code_block(doc, code)

add_heading3(doc, "The Knot Canada")
add_body(doc, "Repeat the exact same steps and copy-paste the same name and description on theknot.com.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 14: LINKEDIN POST
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 14: LinkedIn Post")
add_heading2(doc, "LinkedIn — Copy-Paste Post")

add_body(doc, "Post this from your personal LinkedIn profile for maximum reach.")
add_code_block(doc,
    "I've been quietly building something for the past few months and it's finally live.\n\n"
    "Melaa.ca — a vendor marketplace built specifically for South Asian weddings in the GTA.\n\n"
    "The problem: South Asian couples spend hours searching through generic directories and half the results don't understand the culture, the traditions, or what makes a South Asian wedding what it is.\n\n"
    "Mela fixes that. Every vendor is local to the GTA and familiar with South Asian weddings. Photographers, decorators, caterers, mehndi artists, DJs — all in one place.\n\n"
    "Free to browse. Free for vendors to list.\n\n"
    "If you know anyone planning a South Asian wedding in the GTA, or anyone who works in the wedding industry, I'd be incredibly grateful if you shared this.\n\n"
    "melaa.ca"
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 15: PARTNERSHIP & INFLUENCER OUTREACH
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 15: Partnership & Influencer Outreach")
add_heading2(doc, "Influencer & Blog Partnership Outreach")

add_heading3(doc, "How to Find Influencers")
add_body(doc, "Search these hashtags on Instagram and look for accounts with 5,000–100,000 followers who post about South Asian weddings:")
for tag in ["#southasianweddingblog", "#desibride", "#southasianbride", "#gtabride", "#canadiandesicouple", "#punjabicouple"]:
    add_bullet(doc, tag)

add_heading3(doc, "Influencer DM Script")
add_code_block(doc,
    "Hey [Name]! Your content is incredible — I've been following your work for a while.\n\n"
    "I'm [Your Name], Co-Founder of Melaa.ca — a marketplace specifically for South Asian wedding vendors in the GTA. We're in early days and handpicking who we work with.\n\n"
    "Would you be open to a quick conversation about featuring Mela to your audience? Happy to offer your followers something exclusive."
)

add_heading3(doc, "Blog Outreach Email")
add_body(doc, "Subject:")
add_code_block(doc, "Collaboration — South Asian wedding vendor marketplace for GTA")
add_body(doc, "Body:")
add_code_block(doc,
    "Hi [Name],\n\n"
    "I'm [Your Name], Co-Founder of Mela (melaa.ca) — we've built the GTA's first marketplace exclusively for South Asian wedding vendors.\n\n"
    "I've been a reader of [blog name] for a while and think our audiences are a perfect match. We'd love to:\n\n"
    "- Write a guest post about finding South Asian vendors in the GTA (SEO value for you)\n"
    "- Feature your blog as a resource on Mela for couples\n"
    "- Cross-promote on our Instagram (growing fast)\n\n"
    "Would you be open to a quick 15-minute chat?\n\n"
    "Best,\n"
    "[Your Name]\n"
    "Mela — melaa.ca"
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 16: WEEKLY OPERATING RHYTHM
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 16: Weekly Operating Rhythm")
add_heading2(doc, "Your Weekly Schedule — Once You're Running")

add_table(
    doc,
    headers=["Day", "Key Actions"],
    rows=[
        ["Monday",
         "Check admin dashboard (melaa.ca/admin) — review new signups, leads, agent logs. Follow up all unanswered DMs from last week. Post 2 Instagram pieces."],
        ["Tuesday",
         "Send 10 new vendor WhatsApp or Instagram DM outreach messages. Post 1 Instagram + 1 TikTok. Engage with all comments from previous posts (reply to every comment)."],
        ["Wednesday",
         "Post in 3 Facebook groups. Film 1–2 TikTok videos. Check agent logs at melaa.ca/admin/org to see what's running."],
        ["Thursday",
         "Partnership/influencer outreach — 3 new contacts. Post 2 Instagram pieces. Reply to all vendor questions or DM replies."],
        ["Friday",
         "Weekly review: vendors added, leads generated, revenue. Plan next week's content calendar. Post in Reddit communities."],
        ["Saturday",
         "Post 1–2 Instagram Stories or Reels. Light engagement — reply to comments, like vendor posts."],
        ["Sunday",
         "Rest or light content creation for the week ahead."],
    ],
    col_widths=[1.2, 5.4],
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 17: BRAND VOICE GUIDE
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 17: Brand Voice Guide")
add_heading2(doc, "How Mela Sounds — Brand Voice")

add_heading3(doc, "We Are")
for item in [
    "Warm and community-first",
    "Confident but not corporate",
    "South Asian, from the GTA, for the GTA",
    "Direct and honest — no fluff",
]:
    add_bullet(doc, item)

add_heading3(doc, "We Are Not")
for item in [
    "Salesy or pushy",
    "Generic or corporate",
    "Formal or stiff",
    "Trying to sound like a big tech company",
]:
    add_bullet(doc, item)

add_table(
    doc,
    headers=["Say This", "Not This"],
    rows=[
        ["Built for our community", "Leveraging synergies"],
        ["Every vendor understands your culture", "World-class vendor network"],
        ["Free to browse, no hidden fees", "Transparent pricing model"],
        ["We're from the GTA", "Serving the Greater Toronto Area region"],
        ["Your dream wedding starts here", "Optimizing your vendor discovery journey"],
    ],
    col_widths=[3.0, 3.0],
)

add_heading3(doc, "Tone by Channel")
for item in [
    "Instagram: Warm, visual, community feel — speak like a friend",
    "TikTok: Casual, honest, relatable — it's okay to be a little funny",
    "Facebook: Helpful, informative — like a community member, not a brand",
    "LinkedIn: Professional founder story, growth narrative",
    "WhatsApp and DMs: Personal, direct, human — no copy-paste feeling",
]:
    add_bullet(doc, item)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 18: KEY METRICS
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 18: Key Metrics to Track")
add_heading2(doc, "What Success Looks Like — Weekly KPIs")

add_table(
    doc,
    headers=["Metric", "Week 1 Target", "Week 4 Target", "Week 12 Target"],
    rows=[
        ["Active Vendors", "8", "20", "50"],
        ["MRR", "$795", "$1,200", "$2,000"],
        ["Instagram Followers", "100", "350", "1,000"],
        ["WhatsApp/DM Outreach", "20", "80", "200"],
        ["Facebook Group Posts", "6", "24", "60"],
        ["Vendor Leads Generated", "5", "25", "75"],
        ["TikTok Videos Posted", "3", "12", "30"],
    ],
    col_widths=[2.1, 1.3, 1.3, 1.3],
)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# SECTION 19: KEY LINKS & ACCOUNTS
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "Section 19: Key Links & Accounts")
add_heading2(doc, "Important Links — Get Access from Ishaan")

add_heading3(doc, "Platform URLs")
for item in [
    "Website: melaa.ca",
    "Admin Dashboard: melaa.ca/admin",
    "Agent Monitor: melaa.ca/admin/org",
    "Outreach Pipeline: melaa.ca/admin/outreach",
    "Vendor Signup: melaa.ca/list-your-business",
    "Pricing: melaa.ca/pricing",
]:
    add_bullet(doc, item)

add_heading3(doc, "Social Accounts to Create (use social@melaa.ca)")
for item in [
    "Instagram: @melaa.ca",
    "TikTok: @melaa.ca",
    "Facebook Page: Mela — South Asian Weddings GTA",
    "Pinterest: @melaa.ca",
    "LinkedIn Company Page: Mela",
]:
    add_bullet(doc, item)

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# CLOSING PAGE
# ═══════════════════════════════════════════════════════════════════
add_heading1(doc, "The Agents Have Your Back")

add_body(doc, "While you're doing all of this human work, here's what's running 24/7 automatically in the background:")

for item in [
    "Every free vendor who signs up gets a 7-part email nurture sequence (Day 1, 7, 14, 30, 60, 85, and 90) — automatically sent at the right time",
    "Every vendor who gets a lead but doesn't reply within 1 hour gets an automatic email nudge",
    "Cold outreach emails are being sent automatically to vendor targets every 15 minutes",
    "SEO blog posts are being written and published automatically",
    "Instagram content is being generated and queued automatically",
    "Revenue is being monitored every 30 minutes and logged",
    "Financial reports and equity statements are sent to the team monthly",
]:
    add_bullet(doc, item)

add_spacer(doc, 8)
add_body(doc, "Your job is the human work the agents can't do: showing up, building relationships, creating authentic content, and making people feel like they're part of something real.")

add_spacer(doc, 6)
p_soul = doc.add_paragraph()
run_soul = p_soul.add_run("The agents handle the scale. You handle the soul.")
run_soul.font.name  = BODY_FONT
run_soul.font.size  = Pt(14)
run_soul.font.bold  = True
run_soul.font.color.rgb = ORANGE
p_soul.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_spacer(doc, 20)
add_body(doc, "Let's build something the South Asian community in the GTA is proud of.")

add_spacer(doc, 10)
p_sig = doc.add_paragraph()
run_sig = p_sig.add_run("— Ishaan Saini, Founder  |  melaa.ca  |  hello@melaa.ca")
run_sig.font.name   = BODY_FONT
run_sig.font.size   = Pt(11)
run_sig.font.italic = True
run_sig.font.color.rgb = DARK


# ── save ──────────────────────────────────────────────────────────────────────
output_path = "/Users/ishaansaini/Downloads/mela/Mela_CoFounder_Playbook.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
